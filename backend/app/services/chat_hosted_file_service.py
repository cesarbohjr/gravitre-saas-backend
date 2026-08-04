"""Durable chat artifact hosting for Phase 2 file-reference chips.

Persists generate_document (and related chat outputs) to Supabase Storage and
returns downloadable file refs + preview HTML for the Preview/Code pane.
"""
from __future__ import annotations

import csv
import html
import io
import re
from datetime import datetime, timezone
from typing import Any
from uuid import uuid4

from app.config import Settings
from app.core.logging import get_logger

logger = get_logger(__name__)

_SAFE = re.compile(r"[^a-zA-Z0-9._-]+")
_SIGNED_TTL_SECONDS = 7 * 24 * 3600


def _sanitize(value: str, *, fallback: str = "file") -> str:
    cleaned = _SAFE.sub("_", (value or "").strip()).strip("._")
    return (cleaned or fallback)[:120]


def _slug_title(title: str) -> str:
    return _sanitize(title.lower().replace(" ", "-"), fallback="document")[:80]


def markdown_to_simple_html(title: str, markdown: str) -> str:
    """Lightweight markdown→HTML for Preview pane (no external markdown dep)."""
    lines = (markdown or "").splitlines()
    body: list[str] = []
    in_ul = False
    in_code = False
    code_buf: list[str] = []

    def close_ul() -> None:
        nonlocal in_ul
        if in_ul:
            body.append("</ul>")
            in_ul = False

    for raw in lines:
        line = raw.rstrip()
        if line.strip().startswith("```"):
            if in_code:
                body.append(
                    "<pre><code>"
                    + html.escape("\n".join(code_buf))
                    + "</code></pre>"
                )
                code_buf = []
                in_code = False
            else:
                close_ul()
                in_code = True
            continue
        if in_code:
            code_buf.append(line)
            continue
        if not line.strip():
            close_ul()
            continue
        if line.startswith("# "):
            close_ul()
            body.append(f"<h1>{html.escape(line[2:].strip())}</h1>")
        elif line.startswith("## "):
            close_ul()
            body.append(f"<h2>{html.escape(line[3:].strip())}</h2>")
        elif line.startswith("### "):
            close_ul()
            body.append(f"<h3>{html.escape(line[4:].strip())}</h3>")
        elif line.lstrip().startswith(("- ", "* ")):
            if not in_ul:
                body.append("<ul>")
                in_ul = True
            body.append(f"<li>{html.escape(line.lstrip()[2:].strip())}</li>")
        else:
            close_ul()
            body.append(f"<p>{html.escape(line)}</p>")
    close_ul()
    if in_code and code_buf:
        body.append("<pre><code>" + html.escape("\n".join(code_buf)) + "</code></pre>")

    return (
        "<!DOCTYPE html><html><head><meta charset='utf-8'/>"
        f"<title>{html.escape(title)}</title>"
        "<style>body{font-family:system-ui,sans-serif;max-width:720px;margin:24px auto;"
        "padding:0 16px;line-height:1.5;color:#111}"
        "pre{background:#f4f4f5;padding:12px;overflow:auto;border-radius:8px}"
        "h1,h2,h3{line-height:1.25}</style></head><body>"
        + "\n".join(body)
        + "</body></html>"
    )


def markdown_to_docx_bytes(title: str, markdown: str) -> bytes:
    from docx import Document

    doc = Document()
    doc.add_heading(title or "Document", level=0)
    for raw in (markdown or "").splitlines():
        line = raw.rstrip()
        if not line.strip():
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.lstrip().startswith(("- ", "* ")):
            doc.add_paragraph(line.lstrip()[2:].strip(), style="List Bullet")
        elif line.strip().startswith("```"):
            continue
        else:
            doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def markdown_to_csv_bytes(markdown: str) -> bytes | None:
    """Extract the first GitHub-flavored markdown table into CSV, if present."""
    lines = [ln.rstrip() for ln in (markdown or "").splitlines()]
    table: list[str] = []
    for ln in lines:
        if ln.strip().startswith("|") and ln.strip().endswith("|"):
            table.append(ln.strip())
        elif table:
            break
    if len(table) < 2:
        return None
    # Drop separator row (|---|---|)
    rows: list[list[str]] = []
    for i, ln in enumerate(table):
        cells = [c.strip() for c in ln.strip("|").split("|")]
        if i == 1 and all(re.fullmatch(r":?-{3,}:?", c or "") for c in cells):
            continue
        rows.append(cells)
    if len(rows) < 2:
        return None
    buf = io.StringIO()
    writer = csv.writer(buf)
    writer.writerows(rows)
    return buf.getvalue().encode("utf-8")


def _minimal_pdf_bytes(title: str, text: str) -> bytes:
    """Tiny text-only PDF (no reportlab) for durable download chips."""
    content_lines = [title, ""] + (text or "").splitlines()
    # PDF text operators — escape parentheses/backslashes
    y = 750
    stream_parts = ["BT", "/F1 11 Tf", "14 TL", "50 750 Td"]
    for i, line in enumerate(content_lines[:80]):
        safe = (
            (line or "")
            .replace("\\", "\\\\")
            .replace("(", "\\(")
            .replace(")", "\\)")
            .encode("latin-1", "replace")
            .decode("latin-1")
        )
        if i == 0:
            stream_parts.append(f"({safe[:90]}) Tj")
        else:
            stream_parts.append("T*")
            stream_parts.append(f"({safe[:90]}) Tj")
        y -= 14
        if y < 50:
            break
    stream_parts.append("ET")
    stream = "\n".join(stream_parts).encode("latin-1", "replace")

    objects: list[bytes] = []
    objects.append(b"1 0 obj<< /Type /Catalog /Pages 2 0 R >>endobj\n")
    objects.append(b"2 0 obj<< /Type /Pages /Kids [3 0 R] /Count 1 >>endobj\n")
    objects.append(
        b"3 0 obj<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>endobj\n"
    )
    objects.append(
        f"4 0 obj<< /Length {len(stream)} >>stream\n".encode("ascii")
        + stream
        + b"\nendstream\nendobj\n"
    )
    objects.append(b"5 0 obj<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>endobj\n")

    out = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for obj in objects:
        offsets.append(len(out))
        out.extend(obj)
    xref_pos = len(out)
    out.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    out.extend(b"0000000000 65535 f \n")
    for off in offsets[1:]:
        out.extend(f"{off:010d} 00000 n \n".encode("ascii"))
    out.extend(
        f"trailer<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref_pos}\n%%EOF\n".encode(
            "ascii"
        )
    )
    return bytes(out)


def status_breakdown_chart_html(title: str, breakdown: dict[str, Any]) -> str:
    """Simple SVG bar chart for analytics Preview pane."""
    items = [(str(k), int(v or 0)) for k, v in (breakdown or {}).items()]
    items = sorted(items, key=lambda x: -x[1])[:8]
    if not items:
        return (
            "<!DOCTYPE html><html><body><p>No run status data in the last 7 days.</p></body></html>"
        )
    max_v = max(v for _, v in items) or 1
    bars = []
    x = 40
    for label, value in items:
        h = int((value / max_v) * 120)
        bars.append(
            f"<rect x='{x}' y='{160 - h}' width='36' height='{h}' fill='#059669'/>"
            f"<text x='{x + 18}' y='178' text-anchor='middle' font-size='9'>"
            f"{html.escape(label[:8])}</text>"
            f"<text x='{x + 18}' y='{156 - h}' text-anchor='middle' font-size='10'>"
            f"{value}</text>"
        )
        x += 50
    width = max(320, x + 20)
    return (
        "<!DOCTYPE html><html><head><meta charset='utf-8'/>"
        f"<title>{html.escape(title)}</title></head><body style='font-family:system-ui'>"
        f"<h2>{html.escape(title)}</h2>"
        f"<svg width='{width}' height='200' viewBox='0 0 {width} 200'>"
        + "".join(bars)
        + "</svg></body></html>"
    )


def codeact_preview_html(description: str | None, result: Any, preview: str) -> str | None:
    """Build Preview HTML when CodeAct result is chartable or structured."""
    if isinstance(result, dict) and result and all(
        isinstance(v, (int, float)) and not isinstance(v, bool) for v in result.values()
    ):
        return status_breakdown_chart_html(description or "Transform result", result)
    if isinstance(result, list) and result and all(isinstance(x, (int, float)) for x in result[:20]):
        breakdown = {str(i): int(v) for i, v in enumerate(result[:12])}
        return status_breakdown_chart_html(description or "Transform series", breakdown)
    safe = html.escape((preview or repr(result))[:4000])
    return (
        "<!DOCTYPE html><html><head><meta charset='utf-8'/></head>"
        f"<body style='font-family:ui-monospace,monospace;padding:16px'>"
        f"<h3>{html.escape(description or 'Code transform')}</h3>"
        f"<pre>{safe}</pre></body></html>"
    )


class ChatHostedFileService:
    """Upload chat artifacts and return file-reference chip payloads."""

    def persist_document(
        self,
        client: Any,
        settings: Settings,
        *,
        org_id: str,
        title: str,
        markdown: str,
        environment_name: str = "production",
    ) -> dict[str, Any]:
        slug = _slug_title(title)
        file_id = uuid4().hex[:12]
        stamp = datetime.now(timezone.utc).strftime("%Y%m%d")
        base = "/".join(
            [
                _sanitize(org_id, fallback="org"),
                _sanitize(environment_name, fallback="production"),
                stamp,
                file_id,
            ]
        )
        preview_html = markdown_to_simple_html(title, markdown)
        payloads: list[tuple[str, str, bytes, str]] = [
            (f"{slug}.md", "text/markdown", (markdown or "").encode("utf-8"), "markdown"),
            (f"{slug}.html", "text/html", preview_html.encode("utf-8"), "html"),
            (
                f"{slug}.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                markdown_to_docx_bytes(title, markdown),
                "docx",
            ),
            (f"{slug}.pdf", "application/pdf", _minimal_pdf_bytes(title, markdown), "pdf"),
        ]
        csv_bytes = markdown_to_csv_bytes(markdown)
        if csv_bytes:
            payloads.append((f"{slug}.csv", "text/csv", csv_bytes, "csv"))

        files: list[dict[str, Any]] = []
        for filename, mime, data, role in payloads:
            meta = self._upload(
                client,
                settings,
                path=f"{base}/{filename}",
                data=data,
                content_type=mime,
            )
            download = (meta or {}).get("download_url")
            files.append(
                {
                    "id": f"{file_id}:{role}",
                    "filename": filename,
                    "mime_type": mime,
                    "byte_size": len(data),
                    "role": role,
                    "download_url": download,
                    "storage_path": (meta or {}).get("storage_path"),
                    "storage_bucket": (meta or {}).get("storage_bucket"),
                    "durable": bool(download),
                }
            )

        return {
            "file_id": file_id,
            "hostedFiles": files,
            "previewHtml": preview_html,
            "code": markdown,
            "previewFormat": "markdown",
        }

    def persist_html_chart(
        self,
        client: Any,
        settings: Settings,
        *,
        org_id: str,
        title: str,
        preview_html: str,
        code: str,
        environment_name: str = "production",
    ) -> dict[str, Any]:
        file_id = uuid4().hex[:12]
        stamp = datetime.now(timezone.utc).strftime("%Y%m%d")
        slug = _slug_title(title)
        path = "/".join(
            [
                _sanitize(org_id, fallback="org"),
                _sanitize(environment_name, fallback="production"),
                stamp,
                file_id,
                f"{slug}.html",
            ]
        )
        data = (preview_html or "").encode("utf-8")
        meta = self._upload(client, settings, path=path, data=data, content_type="text/html")
        download = (meta or {}).get("download_url")
        return {
            "file_id": file_id,
            "hostedFiles": [
                {
                    "id": f"{file_id}:html",
                    "filename": f"{slug}.html",
                    "mime_type": "text/html",
                    "byte_size": len(data),
                    "role": "html",
                    "download_url": download,
                    "storage_path": (meta or {}).get("storage_path"),
                    "storage_bucket": (meta or {}).get("storage_bucket"),
                    "durable": bool(download),
                }
            ],
            "previewHtml": preview_html,
            "code": code,
            "previewFormat": "html",
        }

    def _upload(
        self,
        client: Any,
        settings: Settings,
        *,
        path: str,
        data: bytes,
        content_type: str,
    ) -> dict[str, Any] | None:
        if not getattr(settings, "chat_store_hosted_files", True):
            return None
        bucket = (getattr(settings, "chat_artifacts_bucket", None) or "chat-artifacts").strip()
        if not bucket or not data or client is None:
            return None
        try:
            client.storage.from_(bucket).upload(
                path,
                data,
                file_options={"upsert": "true", "content-type": content_type},
            )
        except Exception as exc:  # noqa: BLE001
            logger.warning("chat_hosted_upload_failed path=%s error=%s", path, str(exc)[:200])
            return None

        download_url: str | None = None
        try:
            signed = client.storage.from_(bucket).create_signed_url(path, _SIGNED_TTL_SECONDS)
            if isinstance(signed, dict):
                download_url = (
                    signed.get("signedURL")
                    or signed.get("signedUrl")
                    or signed.get("signed_url")
                )
            elif isinstance(signed, str):
                download_url = signed
        except Exception as exc:  # noqa: BLE001
            logger.warning("chat_hosted_sign_failed path=%s error=%s", path, str(exc)[:200])

        return {
            "storage_bucket": bucket,
            "storage_path": path,
            "storage_bytes": len(data),
            "download_url": download_url,
        }


_service: ChatHostedFileService | None = None


def get_chat_hosted_file_service() -> ChatHostedFileService:
    global _service
    if _service is None:
        _service = ChatHostedFileService()
    return _service

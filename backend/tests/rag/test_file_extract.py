"""Tests for RAG file text extraction (STA-279)."""
from __future__ import annotations

import io
import json

import pytest

from app.rag.file_extract import (
    UnsupportedFileTypeError,
    extract_text_from_bytes,
    ingest_metadata_from_upload,
)


def test_extract_plain_text():
    text = extract_text_from_bytes(b"Hello knowledge base", "notes.txt")
    assert text == "Hello knowledge base"


def test_extract_csv_rows():
    raw = "name,role\nAlice,Admin\nBob,User\n"
    text = extract_text_from_bytes(raw.encode(), "people.csv")
    assert "Alice | Admin" in text
    assert "Bob | User" in text


def test_extract_json_pretty_print():
    payload = {"policy": "retention", "days": 90}
    text = extract_text_from_bytes(json.dumps(payload).encode(), "policy.json")
    assert '"policy"' in text
    assert "90" in text


def test_rejects_unknown_extension():
    with pytest.raises(UnsupportedFileTypeError):
        extract_text_from_bytes(b"binary", "archive.zip")


def test_rejects_empty_file():
    with pytest.raises(ValueError, match="empty"):
        extract_text_from_bytes(b"", "empty.txt")


def test_ingest_metadata_from_upload():
    meta = ingest_metadata_from_upload("guide.md", content_type="text/markdown", extra={"team": "ops"})
    assert meta["ingest_kind"] == "file_upload"
    assert meta["original_filename"] == "guide.md"
    assert meta["team"] == "ops"


def test_extract_docx_text():
    try:
        from docx import Document as DocxDocument
    except ImportError:
        pytest.skip("python-docx not installed")

    buffer = io.BytesIO()
    document = DocxDocument()
    document.add_paragraph("Quarterly playbook")
    document.add_paragraph("Second paragraph")
    document.save(buffer)
    text = extract_text_from_bytes(buffer.getvalue(), "playbook.docx")
    assert "Quarterly playbook" in text
    assert "Second paragraph" in text
